import os
import json
import datetime # Ensure datetime is imported directly
from datetime import datetime as dt_now # Alias for convenience if needed, or just use datetime.datetime
import logging
import traceback # For detailed error logging
import re # For stripping markdown and sanitizing filenames
from flask import Flask, request, url_for, send_from_directory, g, jsonify
import mysql.connector
from mysql.connector import errorcode
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE # Added for style creation

# Import configurations from db_config.py FIRST
try:
    from db_config import (
        DB_CONFIG, UPLOAD_FOLDER, GENERATED_DOCS_DIR, WORD_FILE_BASE_URL,
        IMAGE_OUTPUT_DIR_FLASK
    )
    try:
        from db_config import EXTERNAL_HOSTNAME
    except ImportError:
        EXTERNAL_HOSTNAME = None
except ImportError:
    print("Error: Critical configurations missing from db_config.py. Please ensure it's in the same directory.")
    exit(1)

# Import functions from other project files
from extractWordElement_web import run_extraction # Added for Word parsing
# 解析word文件的教材信息
from word_parser_for_material import parse_word_to_db

import requests # Added for downloading files

app = Flask(__name__)
app.secret_key = os.urandom(24) # Kept for general Flask app session management, though flash is removed

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

APP_HOST_BIND = '0.0.0.0'
APP_PORT = 7777

if EXTERNAL_HOSTNAME and EXTERNAL_HOSTNAME.strip() and EXTERNAL_HOSTNAME.lower() != 'localhost' and EXTERNAL_HOSTNAME != '0.0.0.0':
    app.config['SERVER_NAME'] = f"{EXTERNAL_HOSTNAME.strip()}:{APP_PORT}"
    logger.info(f"SERVER_NAME configured to: {app.config['SERVER_NAME']} for external URLs.")
else:
    default_server_name = f"127.0.0.1:{APP_PORT}" if EXTERNAL_HOSTNAME == '0.0.0.0' else f"localhost:{APP_PORT}"
    app.config['SERVER_NAME'] = default_server_name
    if not EXTERNAL_HOSTNAME or not EXTERNAL_HOSTNAME.strip():
        logger.warning(
            f"EXTERNAL_HOSTNAME is not set or is empty in db_config.py. "
            f"External URLs will default to '{default_server_name}'. "
            f"For production, set EXTERNAL_HOSTNAME in db_config.py to your server's public IP or domain."
        )
    elif EXTERNAL_HOSTNAME.lower() == 'localhost' or EXTERNAL_HOSTNAME == '0.0.0.0':
        logger.warning(
            f"EXTERNAL_HOSTNAME in db_config.py is set to '{EXTERNAL_HOSTNAME}', which is not suitable for external access. "
            f"External URLs will default to '{default_server_name}'. "
            f"For production, use a public IP or domain."
        )

app.config['APPLICATION_ROOT'] = '/'
app.config['PREFERRED_URL_SCHEME'] = 'http'

app.config['DB_CONFIG'] = DB_CONFIG
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['GENERATED_DOCS_DIR'] = GENERATED_DOCS_DIR
app.config['WORD_FILE_BASE_URL'] = WORD_FILE_BASE_URL
app.config['IMAGE_OUTPUT_DIR_FLASK'] = IMAGE_OUTPUT_DIR_FLASK

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['GENERATED_DOCS_DIR'], exist_ok=True)
os.makedirs(app.config['IMAGE_OUTPUT_DIR_FLASK'], exist_ok=True)

def get_db_connection():
    try:
        config_to_use = app.config['DB_CONFIG'].copy()
        conn = mysql.connector.connect(**config_to_use)
        return conn
    except mysql.connector.Error as err:
        logger.error(f"Database connection error: {err}")
        raise

def get_db():
    if 'db_conn' not in g:
        g.db_conn = get_db_connection()
    return g.db_conn

@app.teardown_appcontext
def close_db(error):
    db_conn = g.pop('db_conn', None)
    if db_conn is not None:
        db_conn.close()

# Status keys (English strings) are used internally. The Chinese translation dictionary and context processor are removed.

def update_file_status_in_db(file_id, status_val, message=None, filepath=None): # filepath is for proof_list_filepath
    conn = None
    cursor = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        if filepath: # This is for proof_list_filepath when advice is generated
            sql = "UPDATE file_records SET proof_status = %s, error_message = %s, proof_list_filepath = %s, updated_at = NOW() WHERE id = %s"
            cursor.execute(sql, (status_val, message, filepath, file_id))
        else: # For other status updates, including errors or intermediate steps
            sql = "UPDATE file_records SET proof_status = %s, error_message = %s, updated_at = NOW() WHERE id = %s"
            cursor.execute(sql, (status_val, message, file_id))
        conn.commit()
        logger.info(f"Updated status for {file_id} to {status_val}. Message: {message}, Path: {filepath}")
    except mysql.connector.Error as err:
        logger.error(f"DB Error updating status for {file_id}: {err}")
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

def _setup_document_styles(document):
    """Defines or modifies styles in the Word document."""
    styles = document.styles

    def get_or_add_style(name, style_type):
        try:
            return styles[name]
        except KeyError:
            logger.info(f"Style '{name}' not found, adding new style.")
            return styles.add_style(name, style_type)

    # Normal Style (正文)
    style = get_or_add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = '宋体'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    font.size = Pt(10.5) # 五号
    font.color.rgb = RGBColor(0, 0, 0)
    p_fmt = style.paragraph_format
    p_fmt.line_spacing = Pt(15.75)
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    p_fmt.first_line_indent = Pt(2 * 10.5) 
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW

    # Heading 1
    style = get_or_add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = '黑体'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    font.size = Pt(16) # 三号
    font.color.rgb = RGBColor(0, 0, 0)
    p_fmt = style.paragraph_format
    p_fmt.line_spacing = Pt(27)
    p_fmt.space_before = Pt(3 * 12)
    p_fmt.space_after = Pt(1 * 12)
    p_fmt.first_line_indent = None
    p_fmt.left_indent = None
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2
    style = get_or_add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = '仿宋'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    font.size = Pt(14) # 四号
    font.color.rgb = RGBColor(0, 0, 0)
    p_fmt = style.paragraph_format
    p_fmt.line_spacing = Pt(21)
    p_fmt.space_before = Pt(1 * 12)
    p_fmt.space_after = Pt(1 * 12)
    p_fmt.first_line_indent = None
    p_fmt.left_indent = None
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 3
    style = get_or_add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = '黑体'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    font.size = Pt(12) # 小四
    font.color.rgb = RGBColor(0, 0, 0)
    p_fmt = style.paragraph_format
    p_fmt.line_spacing = Pt(18)
    p_fmt.space_before = Pt(1 * 12)
    p_fmt.space_after = Pt(1 * 12)
    p_fmt.first_line_indent = None
    p_fmt.left_indent = None
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 4
    style = get_or_add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = '宋体'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    font.size = Pt(12) # 小四
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    p_fmt = style.paragraph_format
    p_fmt.line_spacing = Pt(18)
    p_fmt.space_before = Pt(1 * 12)
    p_fmt.space_after = Pt(1 * 12)
    p_fmt.first_line_indent = None
    p_fmt.left_indent = None
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 5
    style = get_or_add_style('Heading 5', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = '黑体'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    font.size = Pt(10.5) # 五号
    font.color.rgb = RGBColor(0, 0, 0)
    p_fmt = style.paragraph_format
    p_fmt.line_spacing = Pt(15.75)
    p_fmt.space_before = Pt(0.5 * 12)
    p_fmt.space_after = Pt(0.5 * 12)
    p_fmt.first_line_indent = None
    p_fmt.left_indent = None
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 6
    style = get_or_add_style('Heading 6', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = '宋体'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    font.size = Pt(10.5) # 五号
    font.color.rgb = RGBColor(0, 0, 0)
    p_fmt = style.paragraph_format
    p_fmt.line_spacing = Pt(15.75)
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    p_fmt.first_line_indent = None
    p_fmt.left_indent = None
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _generate_advice_document_core(file_id_str, config, db_conn_optional=None):
    conn, cursor = None, None
    is_local_conn = False
    original_doc_name_from_db = f"file_{file_id_str}"

    try:
        if db_conn_optional and db_conn_optional.is_connected():
            conn = db_conn_optional
            is_local_conn = False
        else:
            conn = get_db_connection()
            is_local_conn = True

        cursor = conn.cursor(dictionary=True)

        update_file_status_in_db(file_id_str, "processing: fetching content", "CoreLogic: 开始获取内容以生成建议...")
        logger.info(f"CoreLogic Call (Advice Gen): Starting for file_id: {file_id_str}")

        cursor.execute("SELECT original_filename FROM file_records WHERE id = %s", (file_id_str,))
        file_record_info = cursor.fetchone()

        if not file_record_info:
            msg = f"数据库中未找到 file_id {file_id_str} 的记录 (for advice gen)."
            logger.error(f"CoreLogic Call (Advice Gen): {msg}")
            return {"success": False, "message": msg, "output_filename_basename": None, "filepath_for_db": None}

        if not file_record_info['original_filename']:
            msg = f"无法获取 file_id {file_id_str} 的 original_filename (for advice gen)."
            update_file_status_in_db(file_id_str, "error: db record not found", msg)
            logger.error(f"CoreLogic Call (Advice Gen): {msg}")
            return {"success": False, "message": f"错误：无法获取原始文档名 (ID: {file_id_str})，字段缺失。", "output_filename_basename": None, "filepath_for_db": None}

        original_doc_name_from_db = file_record_info['original_filename']
        base_name_no_ext = os.path.splitext(original_doc_name_from_db)[0]
        sanitized_base_name = re.sub(r'[^\w\s\-\u4e00-\u9fff【】]', '_', base_name_no_ext)
        timestamp_str = dt_now.now().strftime('%Y%m%d-%H%M%S')
        output_filename_basename = f"【{sanitized_base_name}】-审校建议清单-{timestamp_str}-{file_id_str}.docx"
        output_filepath_absolute = os.path.join(config['GENERATED_DOCS_DIR'], output_filename_basename)

        cursor.execute(
            "SELECT id, content_id, text_content, page_no, element_type, level "
            "FROM document_contents WHERE file_record_id = %s ORDER BY page_no ASC, sequence_order ASC",
            (file_id_str,)
        )
        contents_from_db_raw = cursor.fetchall()

        if not contents_from_db_raw:
            msg = f"CoreLogic Call (Advice Gen): 未找到ID为 {file_id_str} 的已解析文档内容 (document_contents表为空)."
            update_file_status_in_db(file_id_str, "error: no content", msg)
            logger.error(msg)
            return {"success": False, "message": msg, "output_filename_basename": None, "filepath_for_db": None}

        page_numbers_corrected = []
        last_valid_page_no = None
        temp_page_nos = [row['page_no'] for row in contents_from_db_raw]
        for i, p_no in enumerate(temp_page_nos):
            if p_no is not None and p_no != -1:
                page_numbers_corrected.append(p_no)
                last_valid_page_no = p_no
            else:
                if last_valid_page_no is not None: page_numbers_corrected.append(last_valid_page_no)
                else:
                    forward_valid_page = None
                    for j_idx in range(i + 1, len(temp_page_nos)):
                        if temp_page_nos[j_idx] is not None and temp_page_nos[j_idx] != -1:
                            forward_valid_page = temp_page_nos[j_idx]; break
                    if forward_valid_page is not None:
                        page_numbers_corrected.append(forward_valid_page)
                        last_valid_page_no = forward_valid_page
                    else: page_numbers_corrected.append(None)
        contents_from_db = []
        for i, row in enumerate(contents_from_db_raw):
            new_row = row.copy(); new_row['page_no_corrected'] = page_numbers_corrected[i]; contents_from_db.append(new_row)

        update_file_status_in_db(file_id_str, "processing: fetching chunks", "CoreLogic: 正在获取AI审校数据...")
        cursor.execute( "SELECT id, ai_content FROM document_content_chunks WHERE file_record_id = %s", (file_id_str,) )
        chunks_data = cursor.fetchall()
        all_ai_suggestions_raw, parsing_warnings = [], []
        for chunk_row in chunks_data:
            raw_ai_content, chunk_id_for_log = chunk_row['ai_content'], chunk_row.get('id', 'N/A')
            if raw_ai_content:
                cleaned_ai_content = raw_ai_content.strip()
                if cleaned_ai_content.startswith("```json"): cleaned_ai_content = cleaned_ai_content[len("```json"):].strip()
                if cleaned_ai_content.startswith("```"): cleaned_ai_content = cleaned_ai_content[len("```"):].strip()
                if cleaned_ai_content.endswith("```"): cleaned_ai_content = cleaned_ai_content[:-len("```")].strip()
                try:
                    suggestions_in_chunk = json.loads(cleaned_ai_content, strict=False)
                    if isinstance(suggestions_in_chunk, list): all_ai_suggestions_raw.extend(suggestions_in_chunk)
                    else:
                        warn_msg = f"Parsed ai_content for file {file_id_str}, chunk_id {chunk_id_for_log} is not a list: {type(suggestions_in_chunk)}"; logger.warning(f"CoreLogic Call (Advice Gen): {warn_msg}"); parsing_warnings.append(warn_msg)
                except json.JSONDecodeError as e:
                    warn_msg = f"解析ai_content JSON时出错 for file {file_id_str}, chunk_id {chunk_id_for_log}: {e}. Cleaned: '{cleaned_ai_content[:100]}...'"; logger.warning(f"CoreLogic Call (Advice Gen): {warn_msg}"); parsing_warnings.append(warn_msg)

        if not all_ai_suggestions_raw and not parsing_warnings:
            if not chunks_data or all(not chunk.get('ai_content') for chunk in chunks_data):
                msg = f"CoreLogic Call (Advice Gen): 未找到 file_id {file_id_str} 的AI审校数据 (document_content_chunks 为空或ai_content为空)."
                logger.warning(msg)

        suggestions_map = {}
        for suggestion in all_ai_suggestions_raw:
            if isinstance(suggestion, dict) and "材料id" in suggestion:
                content_id_key = suggestion["材料id"]
                if content_id_key not in suggestions_map: suggestions_map[content_id_key] = []
                suggestions_map[content_id_key].append(suggestion)
            else:
                logger.warning(f"CoreLogic Call (Advice Gen): Invalid suggestion format skipped for file {file_id_str}: {str(suggestion)[:200]}"); parsing_warnings.append(f"Invalid suggestion: {str(suggestion)[:100]}")

        doc = Document()
        doc.core_properties.author = "易审校-V1.0"
        _setup_document_styles(doc)

        main_title_heading = doc.add_heading('审校建议清单', level=0)
        main_title_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_heading(f'原始文档: {original_doc_name_from_db}', level=1)
        p_time = doc.add_paragraph(style='Normal')
        p_time.add_run(f"生成时间: {dt_now.now().strftime('%Y-%m-%d %H:%M:%S')}")
        p_time.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        update_file_status_in_db(file_id_str, "processing: generating doc", "CoreLogic: 正在生成Word建议文档...")

        items_with_at_least_one_suggestion = 0
        any_content_processed_flag = False
        current_displayed_page_no = None

        for i, content_item_db_row in enumerate(contents_from_db):
            full_original_text_from_db = content_item_db_row['text_content'] if content_item_db_row['text_content'] else ""
            content_id_from_doc_contents = content_item_db_row['content_id']
            page_no_to_display = content_item_db_row['page_no_corrected']
            element_type = content_item_db_row.get('element_type')
            heading_level = content_item_db_row.get('level')

            ai_suggestions_for_this_content = suggestions_map.get(content_id_from_doc_contents, [])

            if full_original_text_from_db or ai_suggestions_for_this_content:
                any_content_processed_flag = True

                if page_no_to_display is not None and page_no_to_display != current_displayed_page_no:
                    page_marker_p = doc.add_paragraph()
                    page_marker_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_marker = page_marker_p.add_run(f"【页码：{page_no_to_display}】")
                    current_displayed_page_no = page_no_to_display
                
                paragraph_style_name = 'Normal'
                if element_type == 'heading' and heading_level is not None:
                    if 1 <= heading_level <= 6:
                        paragraph_style_name = f'Heading {heading_level}'
                    else:
                        logger.warning(f"Content ID {content_id_from_doc_contents}: Heading level {heading_level} out of range (1-6). Using 'Normal' style.")
                
                current_paragraph_for_doc = doc.add_paragraph(style=paragraph_style_name)

                if ai_suggestions_for_this_content:
                    items_with_at_least_one_suggestion +=1
                    changes = []
                    for sugg_idx, sugg in enumerate(ai_suggestions_for_this_content):
                        sugg_orig = sugg.get("原始内容")
                        sugg_mod = sugg.get("修改后内容")
                        sugg_reason = sugg.get("出错原因", sugg.get("判断依据"))
                        if sugg_orig:
                            start_index = 0
                            while True:
                                found_pos = full_original_text_from_db.find(sugg_orig, start_index)
                                if found_pos == -1: break
                                changes.append({
                                    "id": f"sugg_{sugg_idx}_{found_pos}", "start": found_pos,
                                    "end": found_pos + len(sugg_orig), "original": sugg_orig,
                                    "modified": sugg_mod if sugg_mod is not None else "",
                                    "reason": sugg_reason
                                })
                                start_index = found_pos + len(sugg_orig)
                    
                    changes.sort(key=lambda x: (x['start'], -(x['end'] - x['start'])))
                    
                    segmented_text_with_ops = []
                    current_pos = 0
                    for change_op in changes:
                        if change_op['start'] > current_pos:
                            segmented_text_with_ops.append({'text': full_original_text_from_db[current_pos:change_op['start']], 'type': 'normal'})
                        
                        segmented_text_with_ops.append({'text': change_op['original'], 'type': 'ai_original', 'change_op': change_op})
                        current_pos = change_op['end']
                    
                    if current_pos < len(full_original_text_from_db):
                        segmented_text_with_ops.append({'text': full_original_text_from_db[current_pos:], 'type': 'normal'})

                    if not changes and full_original_text_from_db:
                         segmented_text_with_ops = [{'text': full_original_text_from_db, 'type': 'normal'}]
                    elif not segmented_text_with_ops and full_original_text_from_db:
                         segmented_text_with_ops = [{'text': full_original_text_from_db, 'type': 'normal'}]

                    for part_info in segmented_text_with_ops:
                        text_part, part_type = part_info['text'], part_info['type']
                        if not text_part: continue

                        if part_type == 'normal':
                            current_paragraph_for_doc.add_run(text_part)
                        elif part_type == 'ai_original':
                            run = current_paragraph_for_doc.add_run(text_part)
                            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                            run.font.strike = True
                            
                            change_op_for_segment = part_info['change_op']
                            if change_op_for_segment.get('modified'):
                                run_mod = current_paragraph_for_doc.add_run(change_op_for_segment['modified'])
                                run_mod.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                            
                            if change_op_for_segment.get('reason'):
                                run_reason = current_paragraph_for_doc.add_run(f" (【出错原因】：{change_op_for_segment['reason']})")
                                run_reason.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                                run_reason.font.size = Pt(9)
                                run_reason.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                
                else:
                    if full_original_text_from_db:
                         current_paragraph_for_doc.add_run(full_original_text_from_db)
                    elif not ai_suggestions_for_this_content :
                         current_paragraph_for_doc.add_run("[此部分无文本内容但可能有格式]").italic = True
            
            if (i + 1) % 100 == 0 and i > 0:
                logger.info(f"CoreLogic Call (Advice Gen): Processed {i+1}/{len(contents_from_db)} for {file_id_str}")

        final_doc_message = None
        if not any_content_processed_flag:
            final_doc_message = "未能从已解析内容中生成任何清单条目。"
            doc.add_paragraph(final_doc_message, style='Normal')
            logger.warning(f"CoreLogic Call (Advice Gen): {final_doc_message} for {file_id_str}")
        elif items_with_at_least_one_suggestion == 0:
            final_doc_message = "文档处理完成，但未找到有效的AI审校建议用于生成清单。"
            doc.add_paragraph(final_doc_message, style='Normal')
            logger.warning(f"CoreLogic Call (Advice Gen): {final_doc_message} for {file_id_str}")
        else:
            logger.info(f"CoreLogic Call (Advice Gen): All {len(contents_from_db)} items processed for {file_id_str}. {items_with_at_least_one_suggestion} had suggestions.")

        doc.save(output_filepath_absolute)
        logger.info(f"CoreLogic Call (Advice Gen): List saved: {output_filepath_absolute}")

        normalized_generated_docs_dir = os.path.normpath(config['GENERATED_DOCS_DIR'])
        base_dir_name_for_db = os.path.basename(normalized_generated_docs_dir)
        filepath_for_db = f"/{base_dir_name_for_db}/{output_filename_basename}"

        status_message_for_db = f"CoreLogic (Advice Gen): 审校清单生成成功."
        if final_doc_message: status_message_for_db += f" {final_doc_message}"
        if parsing_warnings: status_message_for_db += f" 有 {len(parsing_warnings)} 个AI内容解析警告。"

        update_file_status_in_db(file_id_str, "completed: advice generated", status_message_for_db, filepath=filepath_for_db)
        success_message_for_api_or_sse = f"审校清单 '{output_filename_basename}' 已生成。"
        if final_doc_message: success_message_for_api_or_sse += f" 注意: {final_doc_message}"

        if is_local_conn and conn: conn.commit()
        return {"success": True, "message": success_message_for_api_or_sse, "output_filename_basename": output_filename_basename, "filepath_for_db": filepath_for_db }

    except mysql.connector.Error as db_err:
        error_message = f"数据库操作失败 (Advice Gen): {db_err}"; logger.error(f"CoreLogic Call (Advice Gen): DB error for {file_id_str}: {traceback.format_exc()}")
        if file_id_str: update_file_status_in_db(file_id_str, "error: generation failed", f"CoreLogic (Advice Gen): DB Error - {error_message}")
        if is_local_conn and conn and conn.is_connected(): conn.rollback()
        return {"success": False, "message": error_message, "output_filename_basename": None, "filepath_for_db": None}
    except json.JSONDecodeError as json_err:
        error_message = f"审校数据解析失败 (Advice Gen): {json_err}"; logger.error(f"CoreLogic Call (Advice Gen): JSON error for {file_id_str}: {traceback.format_exc()}")
        if file_id_str: update_file_status_in_db(file_id_str, "error: parsing failed", f"CoreLogic (Advice Gen): JSON Parse Error - {error_message}")
        if is_local_conn and conn and conn.is_connected(): conn.rollback()
        return {"success": False, "message": error_message, "output_filename_basename": None, "filepath_for_db": None}
    except Exception as e:
        error_message = f"生成清单时未知错误 (Advice Gen): {type(e).__name__} - {e}"; logger.error(f"CoreLogic Call (Advice Gen): Generic error for {file_id_str}: {traceback.format_exc()}")
        if file_id_str: update_file_status_in_db(file_id_str, "error: generation failed", f"CoreLogic (Advice Gen): Unknown Error - {error_message}")
        if is_local_conn and conn and conn.is_connected(): conn.rollback()
        return {"success": False, "message": error_message, "output_filename_basename": None, "filepath_for_db": None}
    finally:
        if cursor: cursor.close()
        if is_local_conn and conn and conn.is_connected():
            conn.close()
        logger.info(f"CoreLogic Call (Advice Gen): Finished for {file_id_str}.")

# 解析审校需要的word文档
@app.route('/extract_word_element', methods=['POST'])
def extract_word_element_api():
    content_type = request.headers.get('Content-Type')
    if not content_type or 'application/json' not in content_type.lower():
        logger.warning(f"API /extract_word_element: Received incorrect Content-Type: {content_type}")
        return jsonify({"code": 415, "message": "Unsupported Media Type: Content-Type must be application/json"}), 415

    try:
        data = request.get_json()
        if data is None:
            logger.warning("API /extract_word_element: Request body is not valid JSON or is empty.")
            return jsonify({"code": 400, "message": "无效的JSON数据或请求体为空"}), 400
        file_id = data.get('id')
    except Exception as e:
        logger.error(f"API /extract_word_element: Error parsing JSON data: {e}", exc_info=True)
        return jsonify({"code": 400, "message": f"解析JSON数据时出错: {e}"}), 400

    if not file_id:
        return jsonify({"code": 400, "message": "参数 'id' 不能为空 (在JSON中)"}), 400
    if not isinstance(file_id, str) or not file_id.strip():
        return jsonify({"code": 400, "message": "参数 'id' 必须是一个非空的字符串 (在JSON中)"}), 400
    file_id = file_id.strip()

    conn, cursor = None, None
    original_doc_local_path = None
    file_record = None 

    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        cursor.execute("SELECT id, original_filename, filepath FROM file_records WHERE id = %s", (file_id,))
        file_record = cursor.fetchone()

        if not file_record:
            logger.warning(f"API /extract_word_element: File record ID '{file_id}' not found.")
            return jsonify({"code": 404, "message": f"文件记录 ID '{file_id}' 未找到"}), 404

        original_doc_django_path = file_record.get('filepath')
        if not original_doc_django_path:
            msg = f"API /extract_word_element: File record for ID '{file_id}' is missing 'filepath'."
            logger.error(msg)
            update_file_status_in_db(file_id, "error: source filepath missing", "API: 数据库记录中缺少原始文件路径")
            return jsonify({"code": 500, "message": "文件记录缺少源文件路径信息"}), 500
        
        logger.info(f"API /extract_word_element: Cleaning up previous data for file_id: {file_id}")
        cursor.execute("UPDATE file_records SET proof_list_filepath = NULL, error_message = NULL WHERE id = %s", (file_id,))
        cursor.execute("DELETE FROM document_contents WHERE file_record_id = %s", (file_id,))
        cursor.execute("DELETE FROM document_content_chunks WHERE file_record_id = %s", (file_id,))
        conn.commit()

        update_file_status_in_db(file_id, "processing: initializing_parsing", "API: 准备开始解析Word文档...")

        remote_file_url = f"{app.config['WORD_FILE_BASE_URL'].rstrip('/')}/{original_doc_django_path.lstrip('/')}"
        
        original_filename_from_db = file_record.get('original_filename', f'{file_id}_unknown_file')
        base_name_orig, ext_orig = os.path.splitext(original_filename_from_db)
        
        current_ext_to_use = ext_orig
        if not current_ext_to_use:
            _, ext_from_path = os.path.splitext(original_doc_django_path)
            if ext_from_path and ext_from_path.lower() in ['.doc', '.docx', '.rtf']:
                current_ext_to_use = ext_from_path
            else:
                current_ext_to_use = ".docx"
        
        sanitized_base_name = "".join(c if c.isalnum() or c in ('_', '-') else '_' for c in base_name_orig)
        safe_local_filename_with_ext = f"{sanitized_base_name}{current_ext_to_use}"
        
        original_doc_local_path = os.path.join(app.config['UPLOAD_FOLDER'], f"api_dl_{file_id}_{safe_local_filename_with_ext}")

        update_file_status_in_db(file_id, "processing: downloading", f"API: 开始下载原始Word文档: {safe_local_filename_with_ext}")
        logger.info(f"API /extract_word_element: Downloading {remote_file_url} to {original_doc_local_path}")
        
        try:
            response = requests.get(remote_file_url, stream=True, timeout=180)
            response.raise_for_status()
            with open(original_doc_local_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=32768):
                    f.write(chunk)
            logger.info(f"API /extract_word_element: File downloaded successfully: {original_doc_local_path}")
        except requests.exceptions.RequestException as e_req:
            error_msg = f"下载原始Word文件失败: {e_req}"
            logger.error(f"API /extract_word_element: Download failed for {remote_file_url}: {e_req}", exc_info=True)
            update_file_status_in_db(file_id, "error: download failed", f"API DownloadErr: {type(e_req).__name__} - {str(e_req)[:100]}")
            if original_doc_local_path and os.path.exists(original_doc_local_path):
                try: os.remove(original_doc_local_path)
                except OSError: pass 
            return jsonify({"code": 500, "message": error_msg}), 500

        update_file_status_in_db(file_id, "processing: extracting content", "API: 开始解析Word文档内容")
        logger.info(f"API /extract_word_element: Calling run_extraction for {original_doc_local_path}, file_id: {file_id}")
        try:
            run_extraction(original_doc_local_path, file_id, app.config['IMAGE_OUTPUT_DIR_FLASK'])
            
            update_file_status_in_db(file_id, "completed: content extracted", "API: Word文档内容解析与数据库存储完成。")
            logger.info(f"API /extract_word_element: Content extraction successful for file_id: {file_id}")
            return jsonify({"code": 200, "message": "内容提取成功"}), 200
        except Exception as e_extract:
            db_error_message = f"API Extraction Error: {type(e_extract).__name__}: {str(e_extract)[:150]}"
            user_message = f"内容提取失败: {type(e_extract).__name__}"
            
            if hasattr(e_extract, 'args') and isinstance(e_extract.args, tuple) and \
               len(e_extract.args) >= 3 and e_extract.args[2] is not None and \
               isinstance(e_extract.args[2], tuple) and len(e_extract.args[2]) >= 3:
                com_source = e_extract.args[2][1] or "COM Object"
                com_desc = e_extract.args[2][2] or str(e_extract)
                user_message = f"内容提取失败 (COM Error from {com_source}): {com_desc}"
                db_error_message = f"API COM Error ({com_source}): {com_desc[:120]}"
            else:
                user_message = f"内容提取失败: {str(e_extract)}"

            logger.error(f"API /extract_word_element: Content extraction error for file_id {file_id}: {e_extract}", exc_info=True)
            update_file_status_in_db(file_id, "error: extraction failed", db_error_message)
            return jsonify({"code": 500, "message": user_message}), 500

    except mysql.connector.Error as db_err:
        logger.error(f"API /extract_word_element: DB error encountered for file_id {file_id}: {db_err}", exc_info=True)
        if file_id and file_record:
             update_file_status_in_db(file_id, "error: extraction failed", f"API DBError: {str(db_err.msg)[:150]}")
        return jsonify({"code": 500, "message": f"数据库操作错误: {db_err.msg}"}), 500
    except Exception as e_global:
        logger.error(f"API /extract_word_element: Global error for file_id {file_id}: {e_global}", exc_info=True)
        if file_id and file_record:
            update_file_status_in_db(file_id, "error: extraction failed", f"API GlobalError: {type(e_global).__name__} - {str(e_global)[:100]}")
        return jsonify({"code": 500, "message": f"处理时发生未知错误: {e_global}"}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()
        if original_doc_local_path and os.path.exists(original_doc_local_path):
            try:
                os.remove(original_doc_local_path)
                logger.info(f"API /extract_word_element: Cleaned up temporary file: {original_doc_local_path}")
            except OSError as e_remove:
                logger.warning(f"API /extract_word_element: Failed to clean up temporary file {original_doc_local_path}: {e_remove}")

# 解析word格式的教材信息
@app.route('/flattern_word_element', methods=['POST'])
def flattern_word_element_api():
    # 1. --- Validate Request ---
    content_type = request.headers.get('Content-Type')
    if not content_type or 'application/json' not in content_type.lower():
        logger.warning(f"API /flattern_word_element: Received incorrect Content-Type: {content_type}")
        return jsonify({"code": 415, "message": "Unsupported Media Type: Content-Type must be application/json"}), 415

    try:
        data = request.get_json()
        if data is None:
            logger.warning("API /flattern_word_element: Request body is not valid JSON or is empty.")
            return jsonify({"code": 400, "message": "无效的JSON数据或请求体为空"}), 400
        
        doc_id = data.get('id')
        parse_level = data.get('parse_level')
        material_id = data.get('material_id')
        file_path_url = data.get('file_path')

        # Check for missing parameters
        missing_params = []
        if not doc_id: missing_params.append('id')
        if not parse_level: missing_params.append('parse_level')
        if not material_id: missing_params.append('material_id')
        if not file_path_url: missing_params.append('file_path')
        if missing_params:
            return jsonify({"code": 400, "message": f"以下参数不能为空: {', '.join(missing_params)}"}), 400

        # Validate types
        try:
            parse_level = int(parse_level)
            material_id = int(material_id)
            if parse_level <= 0 or parse_level > 8:
                return jsonify({"code": 400, "message": "参数 'parse_level' 必须是 1 到 8 之间的整数"}), 400
        except (ValueError, TypeError):
            return jsonify({"code": 400, "message": "参数 'parse_level' 和 'material_id' 必须是有效的整数"}), 400

    except Exception as e:
        logger.error(f"API /flattern_word_element: Error parsing JSON data: {e}", exc_info=True)
        return jsonify({"code": 400, "message": f"解析JSON数据时出错: {e}"}), 400

    # 2. --- Download File ---
    conn, cursor = None, None
    local_doc_path = None
    try:
        # Sanitize filename from URL
        filename_from_url = os.path.basename(file_path_url)
        safe_filename = re.sub(r'[^\w\.\-]', '_', filename_from_url)
        local_doc_path = os.path.join(app.config['UPLOAD_FOLDER'], f"flattern_{doc_id}_{safe_filename}")

        logger.info(f"API /flattern_word_element: Downloading from {file_path_url} to {local_doc_path}")
        response = requests.get(file_path_url, stream=True, timeout=180)
        response.raise_for_status()
        with open(local_doc_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        logger.info(f"API /flattern_word_element: File downloaded successfully.")

        # 3. --- Process and Insert into DB ---
        conn = get_db_connection()
        cursor = conn.cursor()

        # Clear old data for this material_id to ensure idempotency
        logger.info(f"Clearing previous entries for material_id: {material_id}")
        cursor.execute("DELETE FROM material_contents WHERE material_id = %s", (material_id,))
        logger.info(f"Deleted {cursor.rowcount} old rows for material_id: {material_id}")

        # Call the core parsing logic
        logger.info(f"Starting Word parsing for material_id: {material_id}")
        result = parse_word_to_db(local_doc_path, material_id, parse_level, cursor)

        if result["success"]:
            conn.commit()
            logger.info(f"API /flattern_word_element: Successfully processed and committed for material_id: {material_id}")
            return jsonify({"code": 200, "message": result["message"]}), 200
        else:
            # This case might happen if parsing runs but finds nothing; it's still a success.
            conn.commit() # Commit the deletion of old data
            return jsonify({"code": 200, "message": result.get("message", "处理完成，但未生成任何内容。")}), 200

    except requests.exceptions.RequestException as e:
        error_msg = f"下载Word文件失败: {e}"
        logger.error(f"API /flattern_word_element: Download failed for {file_path_url}: {e}", exc_info=True)
        return jsonify({"code": 500, "message": error_msg}), 500
    
    except Exception as e:
        logger.error(f"API /flattern_word_element: An error occurred for material_id {material_id}: {e}", exc_info=True)
        if conn:
            conn.rollback() # Rollback transaction on any error
            logger.info(f"Database transaction rolled back for material_id: {material_id}")
        return jsonify({"code": 500, "message": f"处理失败: {str(e)}"}), 500
        
    finally:
        # 4. --- Cleanup ---
        if cursor:
            cursor.close()
        if conn:
            conn.close()
        if local_doc_path and os.path.exists(local_doc_path):
            try:
                os.remove(local_doc_path)
                logger.info(f"API /flattern_word_element: Cleaned up temporary file: {local_doc_path}")
            except OSError as e_remove:
                logger.warning(f"API /flattern_word_element: Failed to clean up temp file {local_doc_path}: {e_remove}")
@app.route('/gen_proof_advice', methods=['POST'])
def gen_proof_advice_api():
    content_type = request.headers.get('Content-Type')
    if not content_type or 'application/json' not in content_type.lower():
        logger.warning(f"API /gen_proof_advice: Received incorrect Content-Type: {content_type}")
        return jsonify({"code": 415, "message": "Unsupported Media Type: Content-Type must be application/json"}), 415

    try:
        data = request.get_json()
        if data is None:
             logger.warning("API /gen_proof_advice: Request body is not valid JSON or is empty.")
             return jsonify({"code": 400, "message": "无效的JSON数据或请求体为空"}), 400
        file_id_str = data.get('id')
    except Exception as e:
        logger.error(f"API /gen_proof_advice: Error parsing JSON data: {e}")
        return jsonify({"code": 400, "message": f"解析JSON数据时出错: {e}"}), 400

    if not file_id_str:
        return jsonify({"code": 400, "message": "参数 'id' 不能为空 (在JSON中)"}), 400
    if not isinstance(file_id_str, str) or not file_id_str.strip():
        return jsonify({"code": 400, "message": "参数 'id' 必须是一个非空的字符串 (在JSON中)"}), 400
    file_id_str = file_id_str.strip()

    conn_check, cursor_check = None, None
    try:
        conn_check = get_db_connection()
        cursor_check = conn_check.cursor(dictionary=True)
        cursor_check.execute("SELECT id, proof_status FROM file_records WHERE id = %s", (file_id_str,))
        record = cursor_check.fetchone()
        if not record:
            return jsonify({"code": 404, "message": f"文件记录 ID '{file_id_str}' 未找到"}), 404
    except mysql.connector.Error as db_err:
        logger.error(f"API /gen_proof_advice: DB error checking file_id {file_id_str}: {db_err}")
        return jsonify({"code": 500, "message": "数据库查询错误"}), 500
    finally:
        if cursor_check: cursor_check.close()
        if conn_check: conn_check.close()

    app_config_paths = {'GENERATED_DOCS_DIR': app.config['GENERATED_DOCS_DIR']}
    result = _generate_advice_document_core(file_id_str, app_config_paths)

    if result["success"]:
        with app.app_context():
            # Generates a relative path, e.g., /download_advice_list/some-uuid
            download_link = url_for('download_advice_list', file_id=file_id_str)
        
        return jsonify({
            "code": 200,
            "message": result["message"],
            "file_path": download_link, 
            "filename": result.get("output_filename_basename")
        }), 200
    else:
        return jsonify({"code": 500, "message": result["message"]}), 500

@app.route('/download_advice_list/<string:file_id>')
def download_advice_list(file_id):
    conn, cursor = None, None
    try:
        conn = get_db(); cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT proof_list_filepath, original_filename FROM file_records WHERE id = %s", (file_id,))
        record = cursor.fetchone()
        if record and record['proof_list_filepath']:
            filepath_from_db = record['proof_list_filepath']
            actual_server_directory = app.config['GENERATED_DOCS_DIR']
            filename_component_from_db = os.path.basename(filepath_from_db)
            full_disk_filepath = os.path.join(actual_server_directory, filename_component_from_db)

            if os.path.exists(full_disk_filepath):
                return send_from_directory(
                    directory=actual_server_directory,
                    path=filename_component_from_db,
                    as_attachment=True,
                    download_name=filename_component_from_db
                )
            else:
                logger.error(f"File not found: {full_disk_filepath}. (DB path: {filepath_from_db}, ID: {file_id})")
                update_file_status_in_db(file_id, "error: file path missing", f"清单文件未找到: {filename_component_from_db}")
                return jsonify({"code": 404, "message": "清单文件丢失，请重试生成。"}), 404
        else:
            logger.warning(f"No proof list path for {file_id} or record missing.")
            if record and not record['proof_list_filepath']:
                 update_file_status_in_db(file_id, "error: file path missing", "DB记录中清单文件路径为空")
            return jsonify({"code": 404, "message": "未找到该文件清单记录或路径。"}), 404

    except mysql.connector.Error as err:
        logger.error(f"DB error download for {file_id}: {err}");
        return jsonify({"code": 500, "message": f"下载时数据库错误: {err.msg}"}), 500
    except Exception as e:
        logger.error(f"General error download for {file_id}: {e} - {traceback.format_exc()}");
        return jsonify({"code": 500, "message": f"下载时发生未知错误: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True, host=APP_HOST_BIND, port=APP_PORT)